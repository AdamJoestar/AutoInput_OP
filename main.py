import sys
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
    QLabel, QLineEdit, QPushButton, QMessageBox, 
    QScrollArea, QGridLayout, QTextEdit, QGroupBox, QFileDialog
)
from PyQt5.QtCore import Qt
from docx import Document
from datetime import date
import os
import re 

# --- Konfigurasi File ---
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
# Menggunakan nama file yang diunggah pengguna: New_Template.docx
TEMPLATE_FILENAME = "New_Template.docx" 
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, TEMPLATE_FILENAME)

# --- Definisi Placeholders & Input Fields ---
# Mapping dari Key (untuk kode) ke Placeholder (untuk Word) dan Label (untuk UI)
FIELD_DEFINITIONS = {
    "SAMPLE_DESCRIPTION": {"placeholder": "[TEXT1]", "label": "DESCRIPCIÓN DE LAS MUESTRAS"},
    "DATE_OF_RECEPTION": {"placeholder": "[TEXT2]", "label": "Fecha de Recepción (DD/MM/YYYY)"},
    "COMMERCIAL_BRAND": {"placeholder": "[TEXT4]", "label": "Marca comercial"},
    "MODEL_REFERENCE": {"placeholder": "[TEXT5]", "label": "Referencia del modelo ensayado"},
    "FAMILY": {"placeholder": "[TEXT6]", "label": "Familia"},
    "INSULATION_CLASS": {"placeholder": "[TEXT7]", "label": "Clase de aislamiento"},
    "LIGHT_SOURCE": {"placeholder": "[TEXT8]", "label": "Fuente de luz"},
    "NOMINAL_VOLTAGE": {"placeholder": "[TEXT9]", "label": "Voltaje nominal"},
    "POWER": {"placeholder": "[TEXT10]", "label": "Potencia"},
    "FREQUENCY": {"placeholder": "[TEXT11]", "label": "Frecuencia"},
    "LS_CURRENT_VOLTAGE": {"placeholder": "[TEXT12]", "label": "Corriente/Tensión fuente de luz"},
    "APPLICATION": {"placeholder": "[TEXT13]", "label": "Aplicación"},
    "EXTENSION_MODELS": {"placeholder": "[TEXT14]", "label": "MODELOS DE EXTENSIÓN"},
    "TESTS_PERFORMED": {"placeholder": "[TEXT15]", "label": "ENSAYOS REALIZADOS"},
    "DATE_OF_TEST": {"placeholder": "[TEXT16]", "label": "Fecha de ensayo (DD/MM/YYYY)"},
    "TEST_STANDARDS": {"placeholder": "[TEXT17]", "label": "Normas de ensayo"},
    "CONCLUSIONS": {"placeholder": "[TEXT18]", "label": "CONCLUSIONES"}
}


class DocumentGeneratorApp(QWidget):
    """Aplikasi untuk menginput data dan menghasilkan dokumen Word dari template."""
    def __init__(self):
        super().__init__()
        # Dictionary untuk menyimpan referensi QLineEdit/QTextEdit
        self.input_widgets = {} 
        self.setWindowTitle("Generador de Informes de Pruebas Térmicas")
        self.setStyleSheet("font-size: 14px; font-family: Arial;")
        self.init_ui()

    def init_ui(self):
        """Membangun antarmuka pengguna."""
        main_layout = QVBoxLayout(self)

        # --- Judul ---
        title = QLabel("Ingresar Datos Para el Informe")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 20px; font-weight: bold; margin-bottom: 10px; color: #2C3E50;")
        main_layout.addWidget(title)
        
        # --- Scroll Area untuk banyak input ---
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        content_widget = QWidget()
        form_layout = QVBoxLayout(content_widget)
        form_layout.setSpacing(15)

        # --- Membuat Group Box untuk Input agar terstruktur ---
        
        # Group 1: Product and Sample Information (1-6)
        self.create_input_group(form_layout, "Información del producto y muestras (Campos 1-6)", [
            "SAMPLE_DESCRIPTION", "DATE_OF_RECEPTION", 
            "COMMERCIAL_BRAND", "MODEL_REFERENCE", "FAMILY", "INSULATION_CLASS"
        ])

        # Group 2: Electrical and Application Details (7-12)
        self.create_input_group(form_layout, "Detalles eléctricos y de aplicación (Campos 7-12)", [
            "LIGHT_SOURCE", "NOMINAL_VOLTAGE", "POWER", 
            "FREQUENCY", "LS_CURRENT_VOLTAGE", "APPLICATION"
        ])
        
        # Group 3: Test and Conclusion Details (13-17)
        self.create_input_group(form_layout, "Detalles de Pruebas y Conclusiones (Campos 13-17)", [
            "EXTENSION_MODELS", "TESTS_PERFORMED", "DATE_OF_TEST", 
            "TEST_STANDARDS", "CONCLUSIONS"
        ])


        scroll.setWidget(content_widget)
        main_layout.addWidget(scroll)

        # --- Tombol Generate ---
        self.generate_button = QPushButton("GENERAR DOCUMENTO DE WORD (.docx)")
        self.generate_button.setStyleSheet(
            "background-color: #3498DB; color: white; padding: 12px; border-radius: 8px; font-weight: bold;"
        )
        self.generate_button.clicked.connect(self.generate_document)
        main_layout.addWidget(self.generate_button)

        # --- Informasi Template ---
        info = QLabel(
            f"**Plantilla utilizada:** '{TEMPLATE_FILENAME}'\n"
            f"Asegúrate de que este archivo esté en la carpeta: '{TEMPLATES_DIR}'"
        )
        info.setStyleSheet("font-size: 10px; color: gray; margin-top: 5px;")
        main_layout.addWidget(info)

        self.setLayout(main_layout)
        self.resize(600, 700) # Ukuran awal yang lebih besar untuk banyak input

    def create_input_group(self, parent_layout, title, keys):
        """Membuat group box untuk input yang terorganisir."""
        group_box = QGroupBox(title)
        group_box.setStyleSheet("font-weight: bold; margin-top: 10px;")
        grid_layout = QGridLayout()
        grid_layout.setSpacing(10)
        
        row = 0
        col = 0
        
        for key in keys:
            definition = FIELD_DEFINITIONS[key]
            
            label = QLabel(f"{definition['label']}:")
            
            if key in ["SAMPLE_DESCRIPTION", "EXTENSION_MODELS", "TESTS_PERFORMED", "CONCLUSIONS"]:
                # Gunakan QTextEdit untuk input multi-baris
                input_field = QTextEdit()
                input_field.setMinimumHeight(60)
                grid_layout.addWidget(label, row, 0, 1, 2)
                grid_layout.addWidget(input_field, row + 1, 0, 1, 2)
                row += 2
                col = 0
            else:
                # Gunakan QLineEdit untuk input satu baris
                input_field = QLineEdit()
                input_field.setMinimumHeight(30)

                # Set tanggal hari ini sebagai default untuk field tanggal
            
                # Tata letak 2 kolom
                grid_layout.addWidget(label, row, col)
                grid_layout.addWidget(input_field, row + 1, col)

                col = 1 - col # Pindah ke kolom 1 (atau kembali ke 0)
                if col == 0:
                    row += 2 # Pindah ke baris baru setelah 2 kolom terisi
            
            self.input_widgets[key] = input_field

        group_box.setLayout(grid_layout)
        parent_layout.addWidget(group_box)
        
    def replace_in_paragraph(self, paragraph, placeholder, value):
        """Mengganti placeholder di dalam paragraf, menjaga format teks."""
        if placeholder in paragraph.text:
            # Menggunakan runs untuk menjaga format, tapi lebih kompleks. 
            # Untuk template sederhana, kita bisa langsung ganti text dan biarkan python-docx mengatur format.
            paragraph.text = paragraph.text.replace(placeholder, value)

    def replace_in_tables(self, document, replacement_data):
        """Mengganti placeholder di dalam sel tabel."""
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Ganti di setiap paragraf di dalam sel
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacement_data.items():
                            self.replace_in_paragraph(paragraph, placeholder, value)
                    # Ganti di paragraf yang mungkin ada di dalam shape atau kotak teks (lebih jarang)
                    # Jika ada masalah, fokus pada paragraf standar sudah cukup untuk sebagian besar template.

    def generate_document(self):
        """Logika utama untuk membaca input, memuat template, mengganti placeholder, dan menyimpan file."""
        
        # 1. Kumpulkan semua data input
        replacement_data = {}
        all_required_filled = True
        
        for key, definition in FIELD_DEFINITIONS.items():
            input_widget = self.input_widgets.get(key)
            if isinstance(input_widget, QLineEdit):
                value = input_widget.text().strip()
            elif isinstance(input_widget, QTextEdit):
                value = input_widget.toPlainText().strip()
            else:
                continue

            # Check jika field penting kosong
            if key in ["MODEL_REFERENCE", "COMMERCIAL_BRAND", "TEST_STANDARDS", "CONCLUSIONS"] and not value:
                all_required_filled = False
                QMessageBox.warning(self, "Input Kosong", f"columna obligatoria ('{definition['label']}') no puede estar vacío.")
                return

            replacement_data[definition['placeholder']] = value
        
        if not all_required_filled:
            return

        # 2. Cek keberadaan template
        if not os.path.exists(TEMPLATE_PATH):
            QMessageBox.critical(self, "Error", 
                f"Plantilla no encontrada en: {TEMPLATE_PATH}. "
                f"Por favor coloque el archivo '{TEMPLATE_FILENAME}' que usted proporciona en la carpeta 'templates'."
            )
            return

        # 3. Muat Template
        try:
            document = Document(TEMPLATE_PATH)
        except Exception as e:
            QMessageBox.critical(self, "Error al leer la plantilla", f"Error al cargar la plantilla: {e}")
            return

        # 4. Lakukan Penggantian di Paragraf dan Tabel
        for paragraph in document.paragraphs:
            for placeholder, value in replacement_data.items():
                self.replace_in_paragraph(paragraph, placeholder, value)
        
        self.replace_in_tables(document, replacement_data)

        # 5. Simpan Dokumen Akhir
        try:
            # Gunakan Judul Dokumen dan Tanggal Uji untuk nama file
            judul_dokumen = replacement_data.get(FIELD_DEFINITIONS["MODEL_REFERENCE"]["placeholder"], "Report")
            tanggal_uji = replacement_data.get(FIELD_DEFINITIONS["DATE_OF_TEST"]["placeholder"], date.today().strftime("%d %B %Y"))
            
            # Sanitasi nama file
            safe_judul = judul_dokumen.replace(' ', '_').replace('/', '-').replace('\\', '-')
            safe_tanggal = tanggal_uji.replace(' ', '_').replace('/', '-').replace('\\', '-')
            
            output_filename = f"Generated_Test_Report_{safe_judul}_{safe_tanggal}.docx"

            # Tanyakan ke user di mana menyimpan dan nama file (Save As)
            options = QFileDialog.Options()
            # default name di dialog adalah output_filename
            file_path, _ = QFileDialog.getSaveFileName(self, "Guardar documento como...", output_filename, "Word Documents (*.docx);;All Files (*)", options=options)
            if not file_path:
                # user membatalkan
                QMessageBox.information(self, "Cancelado", "El almacenamiento fue cancelado por el usuario.")
                return
            # Pastikan ekstensi .docx
            if not file_path.lower().endswith('.docx'):
                file_path += '.docx'

            document.save(file_path)
            QMessageBox.information(
                self,
                "¡Listo!",
                f"El documento de Word se creó y se guardó con éxito como:\n{file_path}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error al guardar el archivo", f"Error al guardar el documento: {e}")


if __name__ == '__main__':
    # Pastikan struktur folder template ada sebelum aplikasi berjalan
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
        # Jangan memanggil QMessageBox sebelum QApplication dibuat (akan memicu error)
        # Tampilkan pesan di console agar script tetap aman dijalankan dari terminal
        print(f"La carpeta 'templates' acaba de ser creada. Por favor, coloque el archivo '{TEMPLATE_FILENAME}' Dentro de él, luego vuelve a ejecutar la aplicación.")
        sys.exit() # Keluar agar user dapat menaruh template

    app = QApplication(sys.argv)
    window = DocumentGeneratorApp()
    window.show()
    sys.exit(app.exec_())
