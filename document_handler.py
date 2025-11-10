from docx import Document
from datetime import date
import os
from config import TEMPLATE_PATH
from fields import FIELD_DEFINITIONS

class DocumentHandler:
    """Handles document generation and placeholder replacement."""

    def __init__(self):
        pass

    def replace_in_paragraph(self, paragraph, placeholder, value):
        """Mengganti placeholder di dalam paragraf, menjaga format teks."""
        if placeholder in paragraph.text:
            # Menggunakan runs untuk menjaga format, tapi lebih kompleks.
            # Untuk template sederhana, kita bisa langsung ganti text dan biarkan python-docx mengatur format.
            paragraph.text = paragraph.text.replace(placeholder, value)
            # Atur font untuk semua runs di paragraf
            for run in paragraph.runs:
                run.font.name = "Gordita Light"
                run.font.size = 9 * 12700  # Font size in twips (1/1440 inch), 9pt = 9*12700

    def replace_in_tables(self, document, replacement_data):
        """Mengganti placeholder di dalam sel tabel."""
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Ganti di setiap paragraf di dalam sel
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacement_data.items():
                            self.replace_in_paragraph(paragraph, placeholder, value)
                    # Ganti di paragraf yang mungkin ada di dalam shape atau kotak teks (lebih jarang)
                    # Jika ada masalah, fokus pada paragraf standar sudah cukup untuk sebagian besar template.

    def replace_in_headers(self, document, replacement_data):
        """Mengganti placeholder di dalam header setiap section."""
        for section in document.sections:
            header = section.header
            # Ganti di paragraf header
            for paragraph in header.paragraphs:
                for placeholder, value in replacement_data.items():
                    self.replace_in_paragraph(paragraph, placeholder, value)
            # Ganti di tabel di header jika ada
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacement_data.items():
                                self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_in_footers(self, document, replacement_data):
        """Mengganti placeholder di dalam footer setiap section."""
        for section in document.sections:
            footer = section.footer
            # Ganti di paragraf footer
            for paragraph in footer.paragraphs:
                for placeholder, value in replacement_data.items():
                    self.replace_in_paragraph(paragraph, placeholder, value)
            # Ganti di tabel di footer jika ada
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacement_data.items():
                                self.replace_in_paragraph(paragraph, placeholder, value)

    def generate_document(self, replacement_data, parent_widget):
        """Logika utama untuk membaca input, memuat template, mengganti placeholder, dan menyimpan file."""

        # 2. Cek keberadaan template
        if not os.path.exists(TEMPLATE_PATH):
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(parent_widget, "Error",
                f"Plantilla no encontrada en: {TEMPLATE_PATH}. "
                f"Por favor coloque el archivo 'New_Template.docx' que usted proporciona en la carpeta 'templates'."
            )
            return

        # 3. Muat Template
        try:
            document = Document(TEMPLATE_PATH)
        except Exception as e:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(parent_widget, "Error al leer la plantilla", f"Error al cargar la plantilla: {e}")
            return

        # 4. Lakukan Penggantian di Paragraf, Tabel, dan Header
        for paragraph in document.paragraphs:
            for placeholder, value in replacement_data.items():
                self.replace_in_paragraph(paragraph, placeholder, value)

        self.replace_in_tables(document, replacement_data)
        self.replace_in_headers(document, replacement_data)
        self.replace_in_footers(document, replacement_data)

        # 5. Simpan Dokumen Akhir
        try:
            # Gunakan Judul Dokumen dan Tanggal Uji untuk nama file
            judul_dokumen = replacement_data.get(FIELD_DEFINITIONS["MODEL_REFERENCE"]["placeholder"], "Report")
            tanggal_uji = replacement_data.get(FIELD_DEFINITIONS["DATE_OF_TEST"]["placeholder"], date.today().strftime("%d %B %Y"))

            # Sanitasi nama file
            safe_judul = judul_dokumen.replace(' ', '_').replace('/', '-').replace('\\', '-')
            safe_tanggal = tanggal_uji.replace(' ', '_').replace('/', '-').replace('\\', '-')

            output_filename = f"Generated_Test_Report_{safe_judul}_{safe_tanggal}.docx"

            # Tanyakan ke user di mana menyimpan dan nama file (Save As)
            from PyQt5.QtWidgets import QFileDialog, QMessageBox
            options = QFileDialog.Options()
            # default name di dialog adalah output_filename
            file_path, _ = QFileDialog.getSaveFileName(parent_widget, "Guardar documento como...", output_filename, "Word Documents (*.docx);;All Files (*)", options=options)
            if not file_path:
                # user membatalkan
                QMessageBox.information(parent_widget, "Cancelado", "El almacenamiento fue cancelado por el usuario.")
                return
            # Pastikan ekstensi .docx
            if not file_path.lower().endswith('.docx'):
                file_path += '.docx'

            document.save(file_path)
            QMessageBox.information(
                parent_widget,
                "¡Listo!",
                f"El documento de Word se creó y se guardó con éxito como:\n{file_path}"
            )
        except Exception as e:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(parent_widget, "Error al guardar el archivo", f"Error al guardar el documento: {e}")
